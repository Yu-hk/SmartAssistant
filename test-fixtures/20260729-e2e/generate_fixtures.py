from __future__ import annotations

import csv
import json
import hashlib
from pathlib import Path
from datetime import datetime, timedelta

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak


ROOT = Path(__file__).resolve().parent
KNOWLEDGE = ROOT / "knowledge"
BUSINESS = ROOT / "business-data"
KNOWLEDGE.mkdir(parents=True, exist_ok=True)
BUSINESS.mkdir(parents=True, exist_ok=True)

POLICY_TEXT = """E2E 客服知识库测试资料
批次：E2E-KB-20260729
适用范围：线上商品、订单、库存、物流与售后客服

一、退货与换货政策
1. 普通商品自签收次日起 7 个自然日内可申请无理由退货，商品、附件、包装和赠品应保持完整。
2. 超过 7 日但属于质量问题的，客服应先登记故障现象、订单号和照片；确认后进入质量售后流程，不受 7 日无理由期限限制。
3. 定制商品、已激活的软件授权、明显使用痕迹或缺少关键配件的商品，原则上不支持无理由退货；质量问题仍按质量售后处理。
4. 生鲜商品需在签收后 24 小时内反馈破损、变质或缺斤少两，并提供开箱照片；审核通过后优先补发或退款。
5. 退货运费：非质量原因由买家承担；质量问题、错发或漏发由平台承担，客服应生成售后工单并标记“平台承担运费”。

二、发货与物流时效
1. 现货商品支付成功后 24 小时内出库；库存紧张商品预计 3—5 个工作日发货，客服不得承诺具体小时数。
2. 普通地区配送时效为出库后 1—3 个工作日，偏远地区增加 2—5 个工作日。节假日以公告为准。
3. 物流状态为“运输中”超过 72 小时无轨迹更新时，客服应创建物流催件记录；超过 7 天无更新则升级人工专员。
4. “已签收”不等于售后结束。用户反馈未收到货时，需核验签收证明、收货地址和联系电话，再提交派送调查。

三、客服回答规范
1. 涉及订单、物流、退款和库存的事实必须优先调用业务工具或查询结果，不得凭空猜测。
2. 不能确认时要明确说明缺少的信息，并列出下一步需要用户提供的字段。
3. 回复中应引用订单号、商品编码或物流单号等可核对标识；不得展示完整手机号、地址或支付凭证。
4. 质量问题、食品安全、重复扣款和疑似欺诈属于高风险场景，应转人工并保留 requestId。

四、测试核对点
- 订单 ORD-E2E-0001：已发货，顺丰单号 E2E-SF-0001，最后节点为“杭州分拨中心，运输中”。
- 订单 ORD-E2E-0002：已签收，圆通单号 E2E-YT-0002，签收时间为 2026-07-28 16:20。
- 商品 E2E-PROD-0001：Aurora 无线降噪耳机，标准价 1299 元，杭州仓可售库存 86 件。
- 商品 E2E-PROD-0002：云杉人体工学椅，标准价 1899 元，上海仓可售库存 14 件，锁定库存 6 件。
"""

POLICY_MD = """# E2E 客服知识库：售后与物流决策表

批次：`E2E-KB-20260729`
文档用途：验证 Markdown/TXT/Word/PDF 的解析、分块、元数据和检索一致性。

## 退货判定矩阵

| 场景 | 时限 | 处理 | 费用 |
| --- | --- | --- | --- |
| 普通商品无理由退货 | 签收次日起 7 个自然日 | 商品完整即可申请 | 买家承担 |
| 质量问题 | 发现后尽快登记 | 记录故障并进入质量售后 | 平台承担 |
| 生鲜破损/变质 | 签收后 24 小时 | 提供开箱照片，优先补发或退款 | 平台承担 |
| 定制商品 | 无理由不适用 | 质量问题仍可售后 | 视责任判定 |

## 物流升级规则

- 运输中超过 72 小时没有新轨迹：创建催件记录。
- 超过 7 天没有更新：升级人工专员。
- 已签收但用户称未收到：核验签收证明、地址、联系电话后发起派送调查。

## 关键测试数据

`ORD-E2E-0001` / 顺丰 / `E2E-SF-0001` / 运输中
`ORD-E2E-0002` / 圆通 / `E2E-YT-0002` / 已签收
`E2E-PROD-0001` / Aurora 无线降噪耳机 / 杭州仓可售 86 件
"""

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",14,7),("Heading 3",12,"1F4D78",10,5)]:
        st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.line_spacing = 1.25
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("客服知识库操作手册（E2E 测试版）")
    run.font.name = "Calibri"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = RGBColor.from_string("0B2545")
    meta = doc.add_paragraph("批次 E2E-KB-20260729 · 版本 v1 · 适用于商品/订单/库存/物流客服")
    meta.runs[0].font.color.rgb = RGBColor.from_string("666666")
    doc.add_heading("1. 退货与售后规则", level=1)
    doc.add_paragraph("本节用于测试标题、段落和列表分块。客服在生成回答前应先确认订单状态和责任归属。")
    for text in ["普通商品签收次日起 7 个自然日内可申请无理由退货，商品和附件应完整。", "超过 7 日但确认属于质量问题的，进入质量售后流程，不受无理由期限限制。", "生鲜破损或变质需在签收后 24 小时内反馈并提供开箱照片。", "定制商品无理由退货不适用，但质量问题仍可申请售后。"]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent = Inches(0.375); p.paragraph_format.first_line_indent = Inches(-0.188); p.add_run(text)
    doc.add_heading("2. 物流时效与升级", level=1)
    doc.add_paragraph("现货商品支付成功后 24 小时内出库；运输中超过 72 小时无轨迹更新要创建催件记录，超过 7 天升级人工专员。")
    doc.add_heading("3. E2E 核对字段", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2400, 3000, 3960])
    headers = ["类型", "标识", "预期事实"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value; set_cell_shading(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs: r.font.bold = True
    rows = [
        ("订单", "ORD-E2E-0001", "顺丰 E2E-SF-0001，运输中，最后节点杭州分拨中心"),
        ("订单", "ORD-E2E-0002", "圆通 E2E-YT-0002，已签收，2026-07-28 16:20"),
        ("商品", "E2E-PROD-0001", "Aurora 无线降噪耳机，1299 元，杭州仓可售 86 件"),
        ("商品", "E2E-PROD-0002", "云杉人体工学椅，1899 元，上海仓可售 14 件，锁定 6 件"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row): cell.text = value
    doc.add_paragraph("来源：E2E 测试夹具；真实客服回答必须以业务工具查询结果为准。")
    doc.add_heading("4. 回答安全边界", level=1)
    doc.add_paragraph("不得猜测订单或库存事实；无法确认时说明缺失字段。回复中只展示脱敏后的联系方式和地址，并保留 requestId 供人工追踪。")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("SmartAssistant E2E · 内部测试资料")
    fr.font.size = Pt(9); fr.font.color.rgb = RGBColor.from_string("777777")
    out = KNOWLEDGE / "customer_service_reference.docx"
    doc.save(out)

def build_pdf():
    # simhei.ttf is a TrueType CJK font supported by ReportLab on Windows.
    font_path = r"C:\Windows\Fonts\simhei.ttf"
    pdfmetrics.registerFont(TTFont("NotoSC", font_path))
    out = KNOWLEDGE / "logistics_inventory_playbook.pdf"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="NotoSC", fontSize=9.5, leading=15, spaceAfter=6)
    h1 = ParagraphStyle("H1CN", parent=body, fontSize=15, leading=21, textColor=colors.HexColor("#0B2545"), spaceBefore=8, spaceAfter=8)
    h2 = ParagraphStyle("H2CN", parent=body, fontSize=11.5, leading=17, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=5)
    story = [Paragraph("物流与库存客服核对手册（E2E 测试版）", h1), Paragraph("批次 E2E-KB-20260729 · 用于验证 PDF 解析、表格抽取与检索", body)]
    story += [Paragraph("一、库存字段解释", h2), Paragraph("可售库存 = 在库数量 - 锁定数量 - 质检冻结数量。客服不得把锁定库存当作可售库存；跨仓查询必须同时报告仓库名称。", body)]
    data = [["商品编码", "仓库", "在库", "锁定", "质检冻结", "可售"], ["E2E-PROD-0001", "杭州仓", "100", "12", "2", "86"], ["E2E-PROD-0002", "上海仓", "25", "6", "5", "14"], ["E2E-PROD-0003", "广州仓", "60", "10", "0", "50"]]
    table = Table(data, colWidths=[31*mm, 22*mm, 15*mm, 15*mm, 22*mm, 15*mm], repeatRows=1)
    table.setStyle(TableStyle([("FONTNAME", (0,0), (-1,-1), "NotoSC"), ("FONTSIZE", (0,0), (-1,-1), 8.5), ("LEADING", (0,0), (-1,-1), 12), ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF5")), ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#9AA7B2")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4)]))
    story += [table, Spacer(1, 8), Paragraph("二、物流节点核对", h2), Paragraph("ORD-E2E-0001 当前为运输中，承运商顺丰，运单 E2E-SF-0001，最近节点为杭州分拨中心。ORD-E2E-0002 当前为已签收，承运商圆通，运单 E2E-YT-0002，签收时间 2026-07-28 16:20。", body), Paragraph("三、异常升级", h2), Paragraph("运输中超过 72 小时无更新，先创建催件记录；超过 7 天无更新，升级人工专员。已签收未收到货时，先核验签收证明、地址和电话，再发起派送调查。", body), Paragraph("四、回答示例", h2), Paragraph("用户问：ORD-E2E-0001 到哪了？正确答复应包含订单号、顺丰、运输中、杭州分拨中心，以及下一步催件阈值，不得编造预计送达时间。", body)]
    SimpleDocTemplate(str(out), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm, title="E2E Logistics Inventory Playbook").build(story)

def write_data():
    (KNOWLEDGE / "customer_service_policy.txt").write_text(POLICY_TEXT, encoding="utf-8")
    (KNOWLEDGE / "returns_logistics_matrix.md").write_text(POLICY_MD, encoding="utf-8")
    products = []
    categories = ["耳机", "办公椅", "显示器", "移动电源", "键盘", "咖啡机", "旅行箱", "运动鞋"]
    for i in range(1, 121):
        code = f"E2E-PROD-{i:04d}"
        products.append({"product_code": code, "product_name": f"E2E {categories[(i-1)%len(categories)]} {i:04d}", "price": f"{199 + (i*37)%4800:.2f}", "stock": "充足" if i % 5 else "紧张", "spec": f"测试规格 {i}；批次 E2E-INV-20260729", "color": "黑色/银色"})
    products[0].update({"product_name": "Aurora 无线降噪耳机", "price": "1299.00"})
    products[1].update({"product_name": "云杉人体工学椅", "price": "1899.00"})
    with (BUSINESS / "products.csv").open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=products[0].keys()); w.writeheader(); w.writerows(products)
    with (BUSINESS / "inventory.csv").open("w", newline="", encoding="utf-8-sig") as f:
        fields = ["product_code", "warehouse", "on_hand", "reserved", "quality_hold", "available", "batch_id"]
        w = csv.DictWriter(f, fieldnames=fields); w.writeheader()
        for i, p in enumerate(products, 1):
            on = 40 + (i * 13) % 180; reserved = (i * 3) % 20; hold = i % 7; available = on - reserved - hold
            if i == 1:
                on, reserved, hold, available = 100, 12, 2, 86
            elif i == 2:
                on, reserved, hold, available = 25, 6, 5, 14
            elif i == 3:
                on, reserved, hold, available = 60, 10, 0, 50
            w.writerow({"product_code": p["product_code"], "warehouse": ["杭州仓", "上海仓", "广州仓"][(i-1)%3], "on_hand": on, "reserved": reserved, "quality_hold": hold, "available": available, "batch_id": f"E2E-INV-20260729-{i:04d}"})
    orders = []; logistics = []
    statuses = ["待付款", "待发货", "已发货", "已签收", "已取消", "退款中"]
    carriers = [("顺丰", "SF"), ("圆通", "YT"), ("中通", "ZT"), ("京东物流", "JD")]
    base_time = datetime(2026, 7, 1, 9, 0)
    for i in range(1, 301):
        order_id = f"ORD-E2E-{i:04d}"; status = statuses[(i-1) % len(statuses)]; carrier, prefix = carriers[(i-1)%len(carriers)]
        if i == 1:
            status = "已发货"
        elif i == 2:
            status = "已签收"
        track = f"E2E-{prefix}-{i:04d}" if status in ("已发货", "已签收") else ""
        orders.append({"order_id": order_id, "user_id": 1, "product_name": products[(i-1)%len(products)]["product_name"], "amount": products[(i-1)%len(products)]["price"], "status": status, "carrier": carrier if track else "", "tracking_no": track, "product_type": categories[(i-1)%len(categories)], "contact_name": f"E2E用户{i:04d}", "contact_phone": f"139{10000000+i:08d}", "shipping_address": f"E2E测试地址{i:04d}", "payment_method": "微信支付", "created_at": (base_time + timedelta(hours=i)).isoformat(sep=" "), "updated_at": (base_time + timedelta(hours=i, minutes=30)).isoformat(sep=" ")})
        if track:
            state = "delivered" if status == "已签收" else "in_transit"
            logistics.append({"tracking_no": track, "order_id": order_id, "carrier": carrier, "status": state, "trajectory": json.dumps([{ "time": (base_time + timedelta(hours=i, minutes=20)).isoformat(sep=" "), "location": "杭州分拨中心" if i == 1 else "E2E中转中心", "desc": "已签收" if state == "delivered" else "运输中" }], ensure_ascii=False)})
    for name, rows in [("orders.csv", orders), ("logistics.csv", logistics)]:
        with (BUSINESS / name).open("w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=rows[0].keys()); w.writeheader(); w.writerows(rows)
    summary = {"batch": "E2E-DATA-20260729", "products": len(products), "inventory_rows": len(products), "orders": len(orders), "logistics": len(logistics), "generated_at": datetime.now().isoformat()}
    (BUSINESS / "manifest.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

if __name__ == "__main__":
    write_data(); build_docx(); build_pdf()
    print(json.dumps({"knowledge": [p.name for p in KNOWLEDGE.iterdir()], "business": [p.name for p in BUSINESS.iterdir()]}, ensure_ascii=False))
